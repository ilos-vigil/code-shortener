import click
import os
import re
import requests
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING

MAX_CHAR = 90
MAX_SEMICOLON = 3


def shorten_markup(file_path):
    # 1. Read from file
    old_lines = open(file_path, 'r').readlines()
    new_lines = []
    current_line = ''

    for line in old_lines:
        # 2. Remove space and EOL
        line = line.lstrip(' ')
        line = line.rstrip('\n')
        line = line.rstrip(' ')

        # 3. Ignore empty line
        if line == "":
            continue

        # 4. Check if multi line is possible
        total_char = len(current_line) + len(line)     
        if total_char <= MAX_CHAR:
            if line[-2:] != '/>' and line[-1] != ">":
                current_line += line.rstrip('\n') + ' '
                continue
            else:
                current_line += line
                line = current_line
        else:
            line, current_line = current_line, line  # swap

        # 5. Add if all condition satisfy
        if line == '}':
            new_lines[len(new_lines) - 1] = new_lines[len(new_lines) - 1].rstrip('\n') + ' }\n'
        else:
            new_lines.append(line + '\n')

        # 6. Reset `current_line`
        current_line = ''
    # 7. Check if last current_line is not empty
    if current_line != '':
        new_lines.append(current_line)

    return new_lines


def shorten_semicolon(file_path):
    # 1. Read from file
    old_lines = open(file_path, 'r').readlines()
    new_lines = []
    current_line = ''

    for line in old_lines:
        # 2. Remove space and EOL
        line = line.lstrip(' ')
        line = line.rstrip('\n')
        line = line.rstrip(' ')

        # 3. Ignore empty and single line comment
        if line == "" or line[:2] == '//':
            continue

        # 4. Check whether multi line is possible
        total_char = len(current_line) + len(line)
        current_line_semicolon = len(re.findall(';', current_line))
        if total_char <= MAX_CHAR and current_line_semicolon < MAX_SEMICOLON:
            if line[-1] != ';' and line[-1] != '}':  # indicate multi line is possible
                current_line += line.rstrip('\n') + ' '
                continue
            elif line[-1] == ';':
                current_line += line.rstrip('\n') + ' '
                continue
            elif current_line != '' and line != '}':  # push char '}' with current_line
                current_line += line
                line = current_line
        else:
            line, current_line = current_line, line  # swap variable

        # 5. Add if all condition satisfy
        if line == '}':
            new_lines[len(new_lines) - 1] = new_lines[len(new_lines) - 1].rstrip('\n') + ' }\n'  # ignore MAX_CHAR rule
        else:
            new_lines.append(line + '\n')

        # 6. Reset `current_line`
        current_line = ''

    # 7. Check if last current_line is not empty
    if current_line != '':
        new_lines.append(current_line)

    return new_lines


def shorten_js(file_path):
    # 1. Read from file
    old_code = open(file_path, 'r').read()
    new_code = requests.post('https://javascript-minifier.com/raw', data={'input': old_code}).text
    new_lines = []

    for new_line in new_code.split(';'):
        new_lines.append(f'{new_line};\n')

    return new_lines

def create_shortened_file(file_dict_list):
    for file_dict in file_dict_list:
        file_path = file_dict['path']
        lines = file_dict['lines']

        seperator_position = file_path.rfind('.')
        output_path = file_path[:seperator_position] + '_short' + file_path[seperator_position:]

        output = open(output_path, 'w')
        output.writelines(lines)
        output.close()


def create_docx(file_dict_list, filename, name, id, title):
    document = docx.Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)

    section = document.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    header = section.header
    footer = section.footer
    header.paragraphs[0].text = f'{name} - {id} | {title}'


    for file_dict in file_dict_list:
        file_path = file_dict['path']
        lines = file_dict['lines']

        # path
        p = document.add_paragraph('')
        p.add_run(file_path).bold = True
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # line
        l = document.add_paragraph(lines)
        l.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    document.save(filename)



@click.command()
@click.argument('path', type=click.Path(exists=True))
@click.option('--format', default='semicolon')
@click.option('--max-char', default=90, help='Maximum character in a line.')
@click.option('--max-semicolon', default=3, help='Maximum semicolon in a line.')
@click.option('--to-docx', is_flag=True, help='Generate .docx contains shortened code')
@click.option('--docx-filename', default='shortened.docx')
@click.option('--docx-name', default='I. J.')
@click.option('--docx-id', default='21xyyzzzzz')
@click.option('--docx-title', default='Lorem Ipsum N')
def main(path, format, max_char, max_semicolon, to_docx, docx_filename, docx_name, docx_id, docx_title):
    """Simple program that save your money"""
    # Set var
    global MAX_CHAR
    global MAX_SEMICOLON
    MAX_CHAR = max_char
    MAX_SEMICOLON = max_semicolon
    file_dict_list = []

    # Check if path is file or directory, then add all absolute path
    file_paths = []
    if os.path.isfile(path):
        file_paths.append(os.path.abspath(path))
    else:
        for (dirpath, dirnames, filenames) in os.walk(path):
            for filename in filenames:
                file_paths.append(os.path.abspath(os.path.join(dirpath, filename)))


    # Shorten each file
    for file_path in file_paths:
        try:
            new_lines = []
            if format.lower() == 'semicolon':
                new_lines = shorten_markup(file_path)
            elif format.lower() == 'markup':
                new_lines = shorten_semicolon(file_path)
            elif format.lower() == 'js':
                new_lines = shorten_js(file_path)
            elif format.lower() == 'same':
                new_lines = open(file_path).readlines()

            file_dict_list.append({
                'path': file_path,
                'lines': new_lines
            })
        except Exception as ex:
            click.echo(f'Something went wrong when shortening code at {file_path}')

    if to_docx:
        create_docx(file_dict_list, docx_filename, docx_name, docx_id, docx_title)
    else:
        create_shortened_file(file_dict_list)


if __name__ == '__main__':
    main()
